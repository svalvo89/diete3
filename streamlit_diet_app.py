
import streamlit as st
from docx import Document
from docx.shared import Pt
import io
from datetime import date
import re
from collections import defaultdict

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None  # ensure PyPDF2 in requirements.txt

st.set_page_config(page_title="Generatore Piano Alimentare", page_icon="🥗", layout="centered")

# -------------------- SIDEBAR -------------------- #
st.sidebar.header("Dati paziente")
patient_name = st.sidebar.text_input("Nome paziente", "Nome Cognome")
place = st.sidebar.text_input("Luogo", "Frascati")
selected_date = st.sidebar.date_input("Data del piano", date.today())
header_date = f"{place}, {selected_date.strftime('%d %B %Y')}"

st.sidebar.header("Carica Anamnesi")
pdf_file = st.sidebar.file_uploader("Scheda anamnesi (PDF)", type=["pdf"])

st.sidebar.header("Opzioni manuali")
manual_kcal = st.sidebar.number_input("Target kcal (0 auto)", 0, 5000, 0)
manual_pathologies = st.sidebar.multiselect(
    "Patologie extra",
    ["Diabete", "Ipercolesterolemia", "Ipertensione", "Endometriosi", "Celiachia", "Intolleranza lattosio"]
)

show_free_meal = st.sidebar.checkbox("Mostra pasto libero", True)

st.title("🥗 Generatore diete adattate all'anamnesi")

# ---------- PDF PARSER ---------- #
def extract_first(pattern, text, cast=float, default=None):
    m = re.search(pattern, text)
    if m:
        try:
            return cast(m.group(1).replace(',', '.'))
        except ValueError:
            return default
    return default

def parse_pdf(pdf):
    data = defaultdict(lambda: None)
    if not PdfReader:
        return data
    try:
        reader = PdfReader(pdf)
        text = "\n".join(page.extract_text() or "" for page in reader.pages).lower()
    except Exception as e:
        st.warning(f"Errore PDF: {e}")
        return data

    data["weight"] = extract_first(r"(\d{2,3})\s*kg", text)
    data["height"] = extract_first(r"(\d{3})\s*cm", text)
    data["age"] = extract_first(r"et[àa]\s*(\d{1,2})", text, int)
    if "maschio" in text or re.search(r"\b(m)\b", text):
        data["sex"] = "M"
    elif "femmina" in text or re.search(r"\b(f)\b", text):
        data["sex"] = "F"

    flags = {
        "diab": "diabetes",
        "glicem": "diabetes",
        "ipercolester": "hyperchol",
        "colesterolo": "hyperchol",
        "ipertensione": "hypertension",
        "pressione alta": "hypertension",
        "endometriosi": "endometriosis",
        "celia": "celiac",
        "gluten": "celiac",
        "lattosio": "lactose",
        "intolleranza al lattosio": "lactose"
    }
    data["conditions"] = set()
    for k, flag in flags.items():
        if k in text:
            data["conditions"].add(flag)

    for kw, val in [("vegetarian", "vegetarian"), ("vegano", "vegan"), ("vegan", "vegan")]:
        if kw in text:
            data["diet"] = val

    if "sedent" in text:
        data["activity"] = 1.2
    elif "leggera" in text or "1-2" in text:
        data["activity"] = 1.375
    elif "moderata" in text or "3-5" in text:
        data["activity"] = 1.55
    elif "intensa" in text or "6-7" in text:
        data["activity"] = 1.725
    else:
        data["activity"] = 1.2
    return data

pdf_data = parse_pdf(pdf_file) if pdf_file else defaultdict(lambda: None)
if manual_pathologies:
    pdf_data["conditions"].update([p.lower() for p in manual_pathologies])

# ---------- KCAL ---------- #
def calc_kcal(sex, w, h, age, act):
    if not all([sex, w, h, age]):
        return 2000
    bmr = 10*w + 6.25*h - 5*age + (5 if sex=="M" else -161)
    tdee = bmr*act
    bmi = w/(h/100)**2
    if bmi > 25:
        tdee -= 400
    return int(tdee)

kcal_target = manual_kcal or calc_kcal(
    pdf_data.get("sex"), pdf_data.get("weight"), pdf_data.get("height"), pdf_data.get("age"), pdf_data.get("activity",1.2)
)

st.write("### Dati rilevati")
st.write({k:v for k,v in pdf_data.items() if v})
st.write(f"**Calorie target stimate:** {kcal_target} kcal")

# ---------- GENERA DIETA ---------- #
def portion(base): return int(base*kcal_target/2000)

def generate_plan():
    doc = Document()
    r = doc.add_paragraph().add_run(header_date)
    r.bold = True; r.font.size = Pt(12)
    doc.add_heading(f"Piano alimentare per {patient_name}", level=1)
    doc.add_paragraph(f"Calorie target: {kcal_target} kcal"); doc.add_paragraph()

    cond = pdf_data.get("conditions", set())
    if cond:
        doc.add_paragraph("Condizioni cliniche: "+", ".join(cond).title()); doc.add_paragraph()

    # Integrazione
    doc.add_heading("INTEGRAZIONE", level=2)
    supp = []
    if "hyperchol" in cond: supp += ["Omega‑3 1 g EPA+DHA", "Riso rosso fermentato 10 mg Monacoline"]
    if "diabetes" in cond: supp += ["Cannella estratto 500 mg", "Cromo picolinato 200 µg"]
    if "endometriosis" in cond: supp += ["Curcumina fitosomiale 500 mg 2×/die", "Magnesio supremo 300 mg"]
    if not supp: supp = ["Multivitaminico quotidiano"]
    for s in supp: doc.add_paragraph(s, style="List Bullet")

    def add_sec(title, opts):
        doc.add_heading(title, level=2)
        for i,o in enumerate(opts,1):
            doc.add_paragraph(f"{i}) {o}", style="List Number")
        doc.add_paragraph()

    gf = "celiac" in cond
    low_salt = "hypertension" in cond

    bread_p = portion(50); cereal_p = portion(40)

    # 5 colazioni
    breakfasts = [
        f"{cereal_p} g fiocchi avena{' GF' if gf else ''} con latte di soia, banana, cannella",
        f"Toast {bread_p} g pane{' SG' if gf else ' integrale'} con avocado e uova strapazzate",
        "Yogurt greco 0 % con frutti di bosco e noci",
        "Pancake proteico (farina d'avena, albume) con burro d'arachidi",
        "Smoothie verde (spinaci, ananas, proteine di pisello)"
    ]
    add_sec("Colazione", breakfasts)

    # 3 snack matt
    snack_m = ["Frutta di stagione (es. mela)", "Mandorle 15 g", "Barretta ai cereali homemade senza zuccheri"]
    add_sec("Spuntino mattutino", snack_m)

    # 7 pranzi
    carb_lunch = portion(70)
    lunches = [
        f"Insalata + {carb_lunch} g {'quinoa' if gf else 'farro'} + pollo grigliato",
        "Riso basmati 80 g + tonno naturale + verdure crude",
        "Pasta integrale 70 g con sugo di pomodoro e basilico",
        "Bowl legumi (ceci, fagioli) + insalata croccante + EVO",
        ("Poke salmone + riso sushi 80 g" if not gf else "Poke salmone + riso bianco 80 g") + " + edamame + avocado",
        (f"Cous cous integrale 70 g" if not gf else "Millet 70 g") + " + verdure grigliate + hummus",
        "Frittata albumi con spinaci + pane integrale 40 g"
    ]
    add_sec("Pranzo", lunches)

    # 3 snack pom
    snack_p = ["Yogurt bianco senza zuccheri con nocciole", "Cioccolato fondente 90 % + noci", "Spremuta d'arancia + pistacchi 10 g"]
    add_sec("Spuntino pomeridiano", snack_p)

    # 7 cene
    bread_n = portion(40*0.6)  # 40g base ridotto 40%
    dinners = [
        f"Verdure al vapore + 2 uova + {bread_n} g pane{' SG' if gf else ' integrale'}",
        "Filetto di salmone al forno + asparagi" + (" (no sale)" if low_salt else ""),
        "Burger di ceci + carote al forno + EVO",
        "Zuppa lenticchie rosse + crostini" + (" SG" if gf else " di pane integrale"),
        "Tacos lattuga con tacchino, avocado e pico de gallo",
        f"Pizza casalinga base {'SG' if gf else 'integrale'} con verdure grigliate",
        "Tempeh saltato + bok choi + riso jasmine 60 g"
    ]
    add_sec("Cena", dinners)

    # 3 snack night
    snack_n = [
        "Tisana melissa + 1 quadratino cioccolato fondente",
        "Kefir di cocco 100 ml",
        "Nice‑cream di banana congelata frullata con cacao"
    ]
    add_sec("Spuntino serale", snack_n)

    if show_free_meal:
        doc.add_heading("PASTO LIBERO", level=2)
        doc.add_paragraph(f"Una volta a settimana: {'pizza SG' if gf else 'pizza margherita'} o altro primo a scelta.")

    doc.add_paragraph()
    footer = doc.add_paragraph("Martina Rastelli\nBIOLOGA NUTRIZIONISTA")
    footer.alignment = 2

    buf = io.BytesIO(); doc.save(buf); return buf

# ---------- MAIN ---------- #
if st.button("Genera piano"):
    file_buf = generate_plan()
    st.success("Piano generato!")
    st.download_button("📥 Scarica DOCX", file_buf.getvalue(), file_name=f"Dieta_{patient_name.replace(' ','_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
